"""Convert markdown docs to formatted Word documents."""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "docs"

DEEP_GREEN = RGBColor(0x07, 0x3D, 0x35)
BODY_FONT = "Calibri"
CODE_FONT = "Consolas"

CONVERSIONS = [
    ("nidaamka_mmf_somali.md", "Nidaamka_MMF_Somali.docx"),
    ("DEFENSE_DEMO_GUIDE.md", "Defense_Demo_Guide.docx"),
    ("defense_presentation_somali.md", "Defense_Presentation_Somali.docx"),
]


def set_default_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for level, size in [(1, 20), (2, 16), (3, 13)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = DEEP_GREEN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(8)


def add_formatted_runs(paragraph, text, base_size=11, base_bold=False, code=False):
    if not text:
        return
    font_name = CODE_FONT if code else BODY_FONT
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            run.font.name = font_name
            run.font.size = Pt(base_size)
            run.bold = base_bold
        chunk = match.group(0)
        if chunk.startswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(base_size)
        else:
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = CODE_FONT
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = font_name
        run.font.size = Pt(base_size)
        run.bold = base_bold


def parse_table_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_table_separator(line):
    return bool(re.match(r"^\|[\s\-:|]+\|$", line.strip()))


def add_table(doc, rows):
    if len(rows) < 1:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            add_formatted_runs(p, cell_text, base_size=10, base_bold=(r_idx == 0))
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = DEEP_GREEN
    doc.add_paragraph()


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    text = "\n".join(lines)
    run = p.add_run(text)
    run.font.name = CODE_FONT
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def polish_title(doc):
    first = None
    for p in doc.paragraphs:
        if p.style.name == "Heading 1":
            first = p
            break
    if first:
        first.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in first.runs:
            run.font.size = Pt(24)
            run.font.color.rgb = DEEP_GREEN


def convert(md_text, out_path):
    doc = Document()
    set_default_styles(doc)

    sections = doc.sections[0]
    sections.top_margin = Cm(2.5)
    sections.bottom_margin = Cm(2.5)
    sections.left_margin = Cm(2.5)
    sections.right_margin = Cm(2.5)

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf = []
    table_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            if not is_table_separator(stripped):
                table_buf.append(parse_table_row(stripped))
            i += 1
            if i >= len(lines) or not lines[i].strip().startswith("|"):
                if table_buf:
                    add_table(doc, table_buf)
                    table_buf = []
            continue

        if table_buf:
            add_table(doc, table_buf)
            table_buf = []

        if stripped in ("---", "***", "___"):
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            add_formatted_runs(p, stripped[2:].strip(), base_size=11, base_bold=False)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_runs(p, stripped[2:].strip(), base_size=11)
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_formatted_runs(p, re.sub(r"^\d+\.\s", "", stripped), base_size=11)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            p = doc.add_paragraph()
            add_formatted_runs(p, stripped, base_size=11, base_bold=True)
            i += 1
            continue

        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.italic = True
            run.font.name = BODY_FONT
            run.font.size = Pt(10)
            i += 1
            continue

        p = doc.add_paragraph()
        add_formatted_runs(p, stripped, base_size=11)
        i += 1

    if table_buf:
        add_table(doc, table_buf)

    polish_title(doc)
    doc.save(out_path)
    return out_path


def main():
    targets = CONVERSIONS
    if len(sys.argv) > 1:
        names = sys.argv[1:]
        targets = [(n, Path(n).stem.replace("_", " ").title().replace(" ", "_") + ".docx") for n in names]

    saved = []
    for md_name, docx_name in targets:
        md_path = DOCS / md_name
        out_path = DOCS / docx_name
        if not md_path.exists():
            print(f"Skip (not found): {md_path}")
            continue
        text = md_path.read_text(encoding="utf-8")
        saved.append(convert(text, out_path))
        print(f"Saved: {out_path}")

    if not saved:
        sys.exit(1)


if __name__ == "__main__":
    main()
